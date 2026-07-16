"""ATS-safe docx renderer: single column, standard headings, no tables or
graphics. Hebrew resumes render right-to-left (w:bidi on every paragraph)."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from aijaa.resume.ir import ResumeIR

HEADINGS = {
    "en": {
        "summary": "Summary", "experience": "Experience", "education": "Education",
        "skills": "Skills", "certifications": "Certifications", "links": "Links",
        "present": "Present",
    },
    "he": {
        "summary": "תקציר", "experience": "ניסיון תעסוקתי", "education": "השכלה",
        "skills": "מיומנויות", "certifications": "הסמכות", "links": "קישורים",
        "present": "היום",
    },
}


def _set_rtl(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        bidi = pPr.makeelement(qn("w:bidi"), {})
        pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def render_docx(ir: ResumeIR, path: str) -> str:
    rtl = ir.language == "he"
    h = HEADINGS[ir.language]
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    def para(text: str, bold: bool = False, size: int | None = None, bullet: bool = False):
        p = doc.add_paragraph(style="List Bullet" if bullet else None)
        run = p.add_run(text)
        run.bold = bold
        if size:
            run.font.size = Pt(size)
        if rtl:
            _set_rtl(p)
            run.font.rtl = True
        return p

    para(ir.full_name, bold=True, size=16)
    if ir.contact_line:
        para(ir.contact_line)
    if ir.links:
        para(" | ".join(ir.links))

    if ir.summary:
        para(h["summary"], bold=True, size=12)
        para(ir.summary)

    if ir.experience:
        para(h["experience"], bold=True, size=12)
        for e in ir.experience:
            dates = f"{e.start} – {e.end or h['present']}"
            title_line = f"{e.title} · {e.company}" + (f" · {e.location}" if e.location else "")
            para(f"{title_line}  ({dates})", bold=True)
            for b in e.bullets:
                para(b.text, bullet=True)

    if ir.education:
        para(h["education"], bold=True, size=12)
        for ed in ir.education:
            bits = [ed.degree, ed.field, ed.institution, ed.year]
            para(", ".join(str(b) for b in bits if b))

    if ir.skills:
        para(h["skills"], bold=True, size=12)
        para(", ".join(ir.skills))

    if ir.certifications:
        para(h["certifications"], bold=True, size=12)
        para(", ".join(ir.certifications))

    doc.save(path)
    return path
