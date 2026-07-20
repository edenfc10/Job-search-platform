"""CV document text extraction for user-uploaded resumes."""

from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader

MAX_CV_BYTES = 10 * 1024 * 1024


class CVParseError(ValueError):
    pass


def extract_cv_text(filename: str, data: bytes) -> str:
    if not data:
        raise CVParseError("CV file is empty")
    if len(data) > MAX_CV_BYTES:
        raise CVParseError("CV file is larger than 10MB")

    suffix = Path(filename or "").suffix.lower()
    if suffix == ".pdf":
        text = _extract_pdf(data)
    elif suffix == ".docx":
        text = _extract_docx(data)
    elif suffix in {".txt", ".text", ".md", ".markdown", ".json"}:
        text = _extract_plain_text(data)
    else:
        raise CVParseError("supported CV formats are PDF, DOCX, TXT, MD, and JSON")

    text = _normalize_text(text)
    if not text:
        raise CVParseError("no readable text found in CV file")
    return text


def _extract_pdf(data: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(data))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as e:  # noqa: BLE001
        raise CVParseError("could not read PDF CV") from e


def _extract_docx(data: bytes) -> str:
    try:
        doc = Document(BytesIO(data))
    except Exception as e:  # noqa: BLE001
        raise CVParseError("could not read DOCX CV") from e
    blocks = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                blocks.append(" | ".join(values))
    return "\n".join(blocks)


def _extract_plain_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise CVParseError("could not decode text CV")


def _normalize_text(text: str) -> str:
    lines = [" ".join(line.split()) for line in (text or "").splitlines()]
    return "\n".join(line for line in lines if line).strip()
