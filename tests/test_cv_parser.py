from io import BytesIO

from docx import Document


async def test_parse_text_cv(client):
    files = {"file": ("cv.txt", b"Dana Levi\nSenior Backend Engineer", "text/plain")}
    r = await client.post("/v1/cv/parse", files=files)
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["filename"] == "cv.txt"
    assert "Dana Levi" in body["text"]
    assert body["characters"] == len(body["text"])


async def test_parse_docx_cv(client):
    doc = Document()
    doc.add_heading("Dana Levi", level=1)
    doc.add_paragraph("Senior Backend Engineer with FastAPI experience.")
    buf = BytesIO()
    doc.save(buf)

    files = {
        "file": (
            "cv.docx",
            buf.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }
    r = await client.post("/v1/cv/parse", files=files)
    assert r.status_code == 200, r.text
    assert "Dana Levi" in r.json()["text"]
    assert "FastAPI" in r.json()["text"]


async def test_parse_rejects_unsupported_cv(client):
    files = {"file": ("cv.pages", b"not supported", "application/octet-stream")}
    r = await client.post("/v1/cv/parse", files=files)
    assert r.status_code == 422
    assert "supported CV formats" in r.json()["detail"]


async def test_local_interpret_handles_general_legal_cv(client):
    text = """SAMPLE RESUME
Daniel Cohen
LAWYER · TEL AVIV, ISRAEL · +972 50 123 4567
daniel@example.com

SKILLS
Legal Database
Government Relations
Regulation
Litigation

LANGUAGES
Hebrew
English

PROFILE
Experienced lawyer. Aspires to act as a VP government relations and regulatory.

EMPLOYMENT HISTORY
Regulation Manager, Trade Association
2019 — Present
• Represented industry members before government authorities and Knesset committees

EDUCATION
2010 - 2014 | LLB Bachelor of Laws, Reichman University
"""
    r = await client.post("/v1/cv/interpret", json={"text": text})
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["profile_patch"]["contact"]["full_name"] == "Daniel Cohen"
    assert body["profile_patch"]["work_history"][0]["title"] == "Regulation Manager"
    skills = {item["text"].casefold() for item in body["profile_patch"]["skills"]}
    assert {"regulation", "litigation", "government relations"} <= skills
    assert "aws" not in skills  # "Laws" must not be misread as the AWS acronym.
    assert body["preferences_patch"]["target_titles"][0].startswith("VP government")


async def test_local_interpret_extracts_complete_backend_cv(client):
    text = """Dana Levi
Tel Aviv, Israel | dana@example.com | +972-50-1234567
LinkedIn: https://linkedin.com/in/danalevi
GitHub: https://github.com/danalevi

Summary
Senior Backend Engineer with Python, FastAPI, PostgreSQL, Kubernetes, AWS, Docker, and Redis experience.

Experience
Senior Backend Engineer · CloudWorks · Tel Aviv (2021-03 – Present)
- Reduced API p95 latency by 40% by rearchitecting the caching layer in Python and Redis.
- Led a team of 5 engineers building FastAPI microservices on Kubernetes serving 2M users.

Backend Engineer · DataNest (2018-06 – 2021-02)
- Built ETL pipelines in Python and PostgreSQL processing 500GB daily.

Education
B.Sc., Computer Science, Tel Aviv University, 2018

Skills
Python, FastAPI, PostgreSQL, Kubernetes, AWS, Docker, Redis, ETL

Languages
English, Hebrew

Preferences: Senior Backend Engineer or Backend Engineer, Tel Aviv or Remote,
minimum salary 30000 ILS, Israeli citizen, avoid gambling companies.
"""

    r = await client.post("/v1/cv/interpret", json={"text": text})

    assert r.status_code == 200, r.text
    body = r.json()
    profile = body["profile_patch"]
    preferences = body["preferences_patch"]

    assert profile["contact"]["full_name"] == "Dana Levi"
    assert [job["company"] for job in profile["work_history"]] == ["CloudWorks", "DataNest"]
    assert [job["title"] for job in profile["work_history"]] == [
        "Senior Backend Engineer",
        "Backend Engineer",
    ]
    assert profile["work_history"][0]["start"] == "2021-03"
    assert profile["work_history"][0]["end"] is None
    assert profile["work_history"][1]["start"] == "2018-06"
    assert profile["work_history"][1]["end"] == "2021-02"
    assert len(profile["work_history"][0]["achievements"]) == 2
    assert len(profile["work_history"][1]["achievements"]) == 1
    assert profile["education"] == [
        {
            "institution": "Tel Aviv University",
            "degree": "B.Sc.",
            "field": "Computer Science",
            "year": "2018",
        }
    ]
    assert preferences["target_titles"][:2] == [
        "Senior Backend Engineer",
        "Backend Engineer",
    ]
    assert preferences["min_salary"] == 30000
    assert preferences["currency"] == "ILS"
    assert preferences["work_authorization"] == "Explicitly stated in CV"
    assert preferences["dealbreakers"] == ["gambling"]


async def test_local_interpret_understands_hebrew_section_headings(client):
    text = """דנה לוי
תל אביב, ישראל | dana@example.com | +972-50-1234567

תקציר
מהנדסת Backend בכירה עם ניסיון ב-Python, FastAPI ו-PostgreSQL.

ניסיון תעסוקתי
מהנדסת Backend בכירה · CloudWorks · תל אביב (2021-03 – היום)
- הובילה צוות של 5 מהנדסים בפיתוח שירותי FastAPI.

השכלה
B.Sc., מדעי המחשב, אוניברסיטת תל אביב, 2018

מיומנויות
Python, FastAPI, PostgreSQL

שפות
עברית, אנגלית
"""

    r = await client.post("/v1/cv/interpret", json={"text": text})

    assert r.status_code == 200, r.text
    profile = r.json()["profile_patch"]
    assert profile["contact"]["full_name"] == "דנה לוי"
    assert profile["work_history"][0]["company"] == "CloudWorks"
    assert profile["work_history"][0]["title"] == "מהנדסת Backend בכירה"
    assert profile["work_history"][0]["start"] == "2021-03"
    assert profile["work_history"][0]["end"] is None
    assert profile["education"][0]["institution"] == "אוניברסיטת תל אביב"
    assert profile["education"][0]["field"] == "מדעי המחשב"
    assert profile["languages"] == ["Hebrew", "English"]


async def test_local_interpret_keeps_headerless_company_first_cv_compatible(client):
    text = """Dana Levi
dana@example.com
CloudWorks - Senior Backend Engineer - 2021-03 to Present - Tel Aviv
- Built reliable Python APIs.
"""

    r = await client.post("/v1/cv/interpret", json={"text": text})

    assert r.status_code == 200, r.text
    job = r.json()["profile_patch"]["work_history"][0]
    assert job["company"] == "CloudWorks"
    assert job["title"] == "Senior Backend Engineer"
    assert job["start"] == "2021-03"
    assert job["end"] is None


async def test_local_interpret_requires_text(client):
    r = await client.post("/v1/cv/interpret", json={"text": "  "})
    assert r.status_code == 422
