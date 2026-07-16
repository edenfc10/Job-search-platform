import os

from docx import Document
from tests.conftest import PROFILE_PATCH

from aijaa.core.models import JobRequirements, Keyword, ProfessionalProfile
from aijaa.resume.ats_score import score
from aijaa.resume.guard import check_ir
from aijaa.resume.ir import Bullet, base_ir_from_profile


def _profile() -> ProfessionalProfile:
    return ProfessionalProfile.model_validate({"seeker_id": "x", **PROFILE_PATCH})


def test_guard_passes_verbatim_ir():
    profile = _profile()
    ir = base_ir_from_profile(profile, "en")
    assert check_ir(ir, profile) == []


def test_guard_catches_fabrications():
    profile = _profile()
    ir = base_ir_from_profile(profile, "en")
    # invented metric
    ir.experience[0].bullets.append(
        Bullet(text="Increased revenue by 300% single-handedly", fact_ids=["f1"])
    )
    # unknown fact id
    ir.experience[0].bullets.append(Bullet(text="Did things", fact_ids=["nope"]))
    # invented skill
    ir.skills.append("COBOL")
    violations = check_ir(ir, profile)
    assert len(violations) == 3
    assert any("300" in v for v in violations)
    assert any("unknown fact_ids" in v for v in violations)
    assert any("COBOL" in v for v in violations)


def _requirements() -> JobRequirements:
    return JobRequirements(
        keywords=[
            Keyword(term="python", must_have=True),
            Keyword(term="kubernetes", aliases=["k8s"], must_have=True),
            Keyword(term="fastapi"),
            Keyword(term="terraform"),
        ],
        seniority="senior",
    )


def test_ats_score_deterministic_and_sensible():
    profile = _profile()
    ir = base_ir_from_profile(profile, "en")
    report = score(ir, _requirements(), job_title="Senior Backend Engineer")
    # python, kubernetes, fastapi hit; terraform misses -> keyword 60*(3+3+1)/(3+3+1+1)
    assert report.score == score(ir, _requirements(), "Senior Backend Engineer").score
    hits = {k.term: k.hit for k in report.keywords}
    assert hits["python"] and hits["kubernetes"] and hits["fastapi"]
    assert not hits["terraform"]
    assert 60 <= report.score <= 100

    # dropping skills lowers the score
    poorer = ir.model_copy(deep=True)
    poorer.skills = []
    poorer.summary = ""
    for e in poorer.experience:
        e.bullets = []
    assert score(poorer, _requirements(), "Senior Backend Engineer").score < report.score


async def test_bilingual_rendering(client, tmp_path):
    from tests.conftest import create_complete_seeker

    seeker_id = await create_complete_seeker(client)
    for lang, heading in (("en", "Experience"), ("he", "ניסיון תעסוקתי")):
        r = await client.post(f"/v1/seekers/{seeker_id}/resume", json={"language": lang})
        assert r.status_code == 200, r.text
        doc = r.json()
        assert doc["language"] == lang
        assert os.path.exists(doc["artifacts"]["docx"])
        assert os.path.exists(doc["artifacts"]["txt"])
        with open(doc["artifacts"]["txt"], encoding="utf-8") as f:
            assert heading.upper() in f.read().upper()
        parsed = Document(doc["artifacts"]["docx"])
        assert any("Dana Levi" in p.text for p in parsed.paragraphs)
        if lang == "he":
            # RTL paragraphs carry w:bidi
            from docx.oxml.ns import qn

            assert any(
                p._p.find(qn("w:pPr")) is not None
                and p._p.find(qn("w:pPr")).find(qn("w:bidi")) is not None
                for p in parsed.paragraphs
                if p.text.strip()
            )
